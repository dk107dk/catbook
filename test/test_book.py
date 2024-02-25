from catbook import Book
from catbook import Files
from catbook import Markup
from catbook import Fonts
from docx import Document


def test_two_line_input_bookfile():
    files = Files()
    files.INPUT = "test/config/one-line-contents.bookfile"
    files.FILES = "test/config/texts"
    book = Book(files=files, markup=Markup(), fonts=Fonts(), document=Document())
    cnt = book.create()
    assert cnt == 2
    assert book.files == 1


def test_two_line_input_bookfile_section_count():
    files = Files()
    files.INPUT = "test/config/one-line-contents.bookfile"
    files.FILES = "test/config/texts"
    book = Book(files=files, markup=Markup(), fonts=Fonts(), document=Document())
    book.create()
    metadata = book.metadata
    assert metadata.count == 1
    print(f"metadata: {metadata}")
    assert metadata.SECTIONS[0] is not None
    assert metadata.SECTIONS[0].NAME == "fish"


def test_coda_bookfile():
    files = Files()
    files.INPUT = "test/config/coda.bookfile"
    files.FILES = "test/config/texts"
    book = Book(files=files, markup=Markup(), fonts=Fonts(), document=Document())
    cnt = book.create()
    assert cnt == 3
    assert book.files == 1
