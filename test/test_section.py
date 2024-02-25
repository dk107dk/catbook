from catbook import RegularSection
from catbook import Markup
from catbook import Fonts
from catbook import Builder
from docx import Document


def test_last_line():
    section = RegularSection([], None, None, None, None)

    assert section._last_line(["", "", ""], 2)
    assert section._last_line(["a", "", ""], 0)
    assert section._last_line(["a", "b", "c"], 2)
    assert not section._last_line(["a", "b", "c"], 0)
    assert not section._last_line(["", "", "c"], 0)


def test_page_break():
    section = RegularSection([], None, None, None, None)
    #
    # opening book
    # pg bk: chapter:False, book: True, last_was_break: False, new_section: False, paragraph_count: 0
    c = False
    b = True
    lwb = False
    s = False
    cnt = 0
    assert not section._needs_page_break_before(
        chapter=c, book=b, last_was_break=lwb, new_section=s, paragraph_count=cnt
    )
    #
    # section
    # pg bk: chapter:False, book: False, last_was_break: False, new_section: True, paragraph_count: 12
    c = False
    b = False
    lwb = False
    s = True
    cnt = 12
    assert section._needs_page_break_before(
        chapter=c, book=b, last_was_break=lwb, new_section=s, paragraph_count=cnt
    )
    #
    # chapter
    # pg bk: chapter:True, book: False, last_was_break: False, new_section: False, paragraph_count: 59
    c = True
    b = False
    lwb = False
    s = False
    cnt = 12
    assert section._needs_page_break_before(
        chapter=c, book=b, last_was_break=lwb, new_section=s, paragraph_count=cnt
    )
    #
    # next book
    # pg bk: chapter:False, book: True, last_was_break: False, new_section: False, paragraph_count: 0
    c = False
    b = True
    lwb = False
    s = False
    cnt = 10
    assert section._needs_page_break_before(
        chapter=c, book=b, last_was_break=lwb, new_section=s, paragraph_count=cnt
    )


def test_handle_block():
    fonts = Fonts()
    markup = Markup()
    markup.BLOCK = "|"
    document = Document()
    section = RegularSection(
        lines=[], markup=markup, fonts=fonts, document=document, metadata=None
    )

    assert section._handle_block("| a block line", 1, 100)
    assert section._handle_block("|another block line", 2, 100)
    assert len(section._block) == 2
    assert not section._handle_block("a non-block line", 1, 100)
    assert section._block is None


def test_handle_quote():
    fonts = Fonts()
    markup = Markup()
    markup.QUOTED_LINE = '"'
    document = Document()
    section = RegularSection(
        lines=[], markup=markup, fonts=fonts, document=document, metadata=None
    )

    assert section._handle_quote('"a quote line', 1, 100)
    assert len(section._quote) == 1
    assert not section._handle_quote("a non-quote line", 2, 100)
    assert section._quote is None


def test_tokenize_example_strings():
    fonts = Fonts()
    markup = Markup()
    markup.QUOTED_LINE = '"'
    document = Document()
    section = RegularSection(
        lines=[], markup=markup, fonts=fonts, document=document, metadata=None
    )

    test = "this is a test"
    words = section._get_words(test)
    print(f"\nwords: {words}")
    n = section._count_words(test)
    assert n == 4

    test = "this's a test"
    words = section._get_words(test)
    print(f"words: {words}")
    n = section._count_words(test)
    assert n == 3

    test = "this-y is a test"
    words = section._get_words(test)
    print(f"words: {words}")
    n = section._count_words(test)
    assert n == 4

    test = "this-y is a test."
    words = section._get_words(test)
    print(f"words: {words}")
    n = section._count_words(test)
    assert n == 4

    test = "this-y. is be a test "
    words = section._get_words(test)
    print(f"words: {words}")
    n = section._count_words(test)
    assert n == 5

    test = "this-y. is''be a test "
    words = section._get_words(test)
    print(f"words: {words}")
    n = section._count_words(test)
    assert n == 5

    test = ".this-y. is''be a test.'"
    words = section._get_words(test)
    print(f"words: {words}")
    n = section._count_words(test)
    assert n == 5


def test_for_spaces_txt():
    builder = Builder()
    builder.init()
    builder.files.OUTPUT = "./test.docx"
    builder.files.INPUT = "test/config/coda.bookfile"
    builder.files.FILES = "test/config/texts"
    builder.build()

    words = builder.book.metadata.words()
    for word in words:
        assert " " not in word
