import csv
import os
import sys
import re
import operator
import traceback
from os import listdir, mkdir
from os.path import isfile, join, exists
from typing import List, Optional
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

import configparser

#
# these are all the types of markups allowed
# note that an empty QUOTED_LINE becomes a blank line.
#
JUMP = "***"
CHAPTER_TITLE = "~"
BOOK_TITLE = "~~"
NEW_SECTION = ">"
BLOCK = "|"
QUOTED_LINE = '"'
WORD_HIGHLIGHT = "|"


class Build(object):
    def __init__(self, contents: str = "contents.csv"):
        self.CONTENTS = contents
        self.OUT = "./out/out.docx"
        self.lines = 0
        self.total_words = 0
        self.words = 0
        self.files = 0
        self.stopped_count = 0
        self.dictionary = dict()
        self.stop = list()
        self._load_stop()
        self._block = None
        self._quote = None
        self._part_break = False
        self._config = configparser.ConfigParser()
        self._config_name = "config.ini"
        self._text_root = "../texts"
        self._font_quote = "Times New Roman"
        self._font_block = "Courier New"
        self._font_title = "Times New Roman"
        self._font = "Times New Roman"
        self._last_was_break = False
        self._linelens = dict()
        self._linelen = False  # no output line lengths

    def load_config(self, name: str, section: str):
        self._load_config(name, section)

    def _load_config(self, name: str = None, section: str = "DEFAULT"):
        if name:
            self._config_name = name
        self._config.read(self._config_name)
        self._text_root = self._config[section]["text_root"]
        self._font_quote = self._config[section]["font_quote"]
        self._font_block = self._config[section]["font_block"]
        self._font_title = self._config[section]["font_title"]
        self._font = self._config[section]["font"]
        self.set_contents(f"{section}.csv")

    def _load_stop(self):
        with open("stop.txt") as s:
            for line in s.readlines():
                self.stop.append(line.strip())

    def print_stats(self):
        print("")
        print(f"total words:  {self.total_words}")
        print(f"stopped words:  {self.words}")
        dl = len(self.dictionary)
        print(f"unique words:  {dl}")
        i = 1
        for k, v in sorted(
            self.dictionary.items(), key=operator.itemgetter(1), reverse=True
        ):
            i = i + 1
            if i < 10:
                print(f"  {dl-i}  {k}: {v}")
        print(f"files:  {self.files}")
        print(f"lines:  {self.lines}")
        print(f"stop count: {self.stopped_count}")

    def set_contents(self, file):
        self.CONTENTS = file
        self.OUT = f"./out/{file.split('.')[0]}.docx"
        try:
            mkdir("./out")
        except:
            pass

    def do(self):
        print("starting build")
        self._clean()
        self._setup()
        if self._validate():
            self._do()
            self._save()

    def _validate(self):
        valid = exists(self.CONTENTS)
        if not valid:
            print(f"no such file: {self.CONTENTS}")
        return valid

    def _save(self):
        print(f"saving {self.OUT}")
        if "/" in self.OUT:
            try:
                d = self.OUT[0 : self.OUT.rindex("/")]
                mkdir(d)
            except:
                pass
        self.doc.save(self.OUT)

    def _setup(self):
        document = Document()
        self.doc = document

    def _clean(self):
        print(f"deleting {self.OUT}")
        try:
            os.remove(self.OUT)
        except Exception as e:
            print(f"clean: {e}")

    def _do(self):
        print(f"reading structure from {self.CONTENTS}")
        with open(self.CONTENTS) as csv_file:
            csv_reader = csv.reader(csv_file, delimiter=",")
            rows = 0
            for row in csv_reader:
                if len(row) == 0:
                    continue
                elif len(row[0]) > 0 and row[0][0] == "#":
                    continue
                else:
                    path = None
                    try:
                        path = f"{row[0]}/{row[1]}"
                        self._append(path)
                        self.files = self.files + 1
                    except Exception as e:
                        print(f"Error: {row}: {rows}: {path}: {e}")
                        print(traceback.format_exc())
                rows = rows + 1

    def _append(self, path: str):
        path = f"{self._text_root}/{path}"
        content = None
        with open(path, "r") as contents:
            lines = contents.readlines()
            line_number = 0
            for line in lines:
                try:
                    self._append_line(lines, line, line_number)
                    line_number = line_number + 1
                    self.lines = self.lines + 1
                except Exception as e:
                    print(f"Error: {line}: {line_number}: {path}: {e}")

    def _is_quote(self, lines: List, line: str, line_number: int) -> bool:
        ll = len(line)
        if ll == 0:
            return False
        if line[0] != QUOTED_LINE:
            return False
        if line_number >= len(lines):
            return False
        if self._quote is None:
            self._quote = []
        self._quote.append(line)
        return True

    def _append_quote(self, lines, line: str, line_number: int):
        i = len(self._quote)
        for aline in self._quote:
            p = self.doc.add_paragraph()
            paragraph_format = p.paragraph_format
            run = p.add_run(f"   {aline[1:]}")
            run.font.name = self._font_quote  #'Times New Roman'
            run.italic = True
            i = i - 1
            if i > 0:
                paragraph_format.space_after = Pt(1)
        self._quote = None

    def _is_block(self, lines: List, line: str, line_number: int) -> Optional[bool]:
        ll = len(line)
        if ll == 0:
            return False
        if line[0] == BLOCK and ll > 1 and line[1] == BLOCK:
            return None
        if line[0] != BLOCK:
            return False
        if line_number >= len(lines):
            return False
        if self._block is None:
            self._block = []
        self._block.append(line)
        return True

    def _append_block(self, lines, line: str, line_number: int):
        i = len(self._block)
        for aline in self._block:
            p = self.doc.add_paragraph()
            paragraph_format = p.paragraph_format
            thisline = aline[1:]
            self._add_to_dictionary(thisline)
            run = p.add_run(f"{thisline}")
            run.italic = True
            run.font.name = self._font_block  # 'Courier New'
            run.font.size = Pt(10)
            i = i - 1
            if i > 0:
                paragraph_format.space_after = Pt(1)
        self._block = None

    def _append_output(self, lines, line: str, line_number: int):
        if self._block is not None:
            # write the block
            self._append_block(lines, line, line_number)
        if self._quote is not None:
            # write the quote
            self._append_quote(lines, line, line_number)
        if self._part_break and self._last_line(lines, line_number):
            p = self.doc.add_paragraph()
            run = p.add_run("")
            run.font.name = self._font  #'Times New Roman'
            run.add_break(WD_BREAK.PAGE)
            self._part_break = False
            self._book_break = False
            self._last_was_break = True

    def _last_line(self, lines: List[str], line_number: int):
        n = len(lines)
        if line_number == n:
            return True
        for r in range(line_number + 1, n):
            if lines[r].strip() != "":
                return False
        return True

    def _append_title(self, lines, line: str, line_number: int):
        self._part_break = len(line) >= 1 and line[0:1] == CHAPTER_TITLE
        self._book_break = len(line) >= 2 and line[0:2] == BOOK_TITLE
        simple_separator = len(line) >= 3 and line[0:3] == JUMP
        close_part = len(line) >= 1 and line[0:1] == NEW_SECTION
        #
        # if we're a _part_break or book_break we need a page break
        #   except if the last file was also a break. in that case
        #   we don't need another blank page
        #
        if (self._part_break and not self._last_was_break) or close_part:
            line = line[2 if self._book_break else 1 :]
            p = self.doc.add_paragraph()
            run = p.add_run("")
            run.font.name = self._font_title  # 'Times New Roman'
            run.add_break(WD_BREAK.PAGE)
        elif self._part_break:
            line = line[2 if self._book_break else 1 :]
        p = None
        if simple_separator:
            #
            # if we're not a _part_break we might be a simple jump separator
            #
            p = self.doc.add_paragraph()
            paragraph_format = p.paragraph_format
            paragraph_format.space_before = Pt(24)
            paragraph_format.space_after = Pt(25)
            run = p.add_run("*                   *                   *")
        else:
            #
            # we're not a simple separator, therefore we must be a heading of some kind
            #
            if self._part_break:
                p = self.doc.add_heading("", 1 if self._book_break else 2)
            else:
                p = self.doc.add_heading("", 3)
                paragraph_format = p.paragraph_format
                paragraph_format.space_before = Pt(30)
                paragraph_format.space_after = Pt(10)
        if self._part_break or simple_separator:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not simple_separator:
            self._add_to_dictionary(line)
            run = p.add_run(line)
            run.font.name = self._font  # 'Times New Roman'
            run.font.color.rgb = RGBColor.from_string("000000")

    def _append_line(self, lines, line: str, line_number: int):
        try:
            line = line.strip()
            #
            # blank line
            #
            if line == "":
                self._append_output(lines, line, line_number)
                return
            #
            # titles
            #
            if line_number == 0:
                return self._append_title(lines, line, line_number)
            #
            # blocks
            #
            block = self._is_block(lines, line, line_number)
            if block is None:
                # found an escaped pipe: || meaning an italicized word starting a line
                line = line[1:]
            elif block:
                # we packed the line into self.block
                return
            elif self._block is not None:
                # write the block
                self._append_block(lines, line, line_number)
            #
            # quotes
            #
            quote = self._is_quote(lines, line, line_number)
            if quote:
                # we packed the line into self.quote
                return
            elif self._quote is not None:
                # write the quote
                self._append_quote(lines, line, line_number)
            #
            # regular line
            #
            self._capture_line_length(line)
            if WORD_HIGHLIGHT in line:
                # check for italics
                block = 0
                p = self.doc.add_paragraph()
                while WORD_HIGHLIGHT in line:
                    start = line.index(WORD_HIGHLIGHT)
                    end = line.index(WORD_HIGHLIGHT, start + 1)
                    front = line[0:start]
                    mid = line[start + 1 : end]
                    back = line[end + 1 :]
                    line = back
                    run = p.add_run("   " if block == 0 else "")
                    run.font.name = self._font  # 'Times New Roman'
                    run = p.add_run(front)
                    run.font.name = self._font  # 'Times New Roman'
                    run = p.add_run(mid)
                    run.italic = True
                    run.font.name = self._font  # 'Times New Roman'
                    if not WORD_HIGHLIGHT in line:
                        run = p.add_run(back)
                        run.font.name = self._font  # 'Times New Roman'
                    block = block + 1
            else:
                p = self.doc.add_paragraph()
                run = p.add_run(f"   {line}")
                run.font.name = self._font  # 'Times New Roman'
            self._last_was_break = False
            self._add_to_dictionary(line)
        except Exception as e:
            print(f"_append_line: error: {e}")
            traceback.print_exc()

    def _capture_line_length(self, line: str) -> None:
        if self._linelen:
            linelen = len(line)
            linelen = linelen / 10
            l = self._linelens.get(f"{linelen}")
            if l == None:
                l = 1
            else:
                l = l + 1
            self._linelens[f"{linelen}"] = l
            if linelen > 40:
                print(f"\nlinelen: {linelen} = {l}")
                print(line)

    def _add_to_dictionary(self, line: str):
        if line and len(line) > 0:
            words = re.split("[ ;\.\-',\?\$%#@!~\"<>/1234567890_\+=]", line)
            for word in words:
                w = f"{word.lower()}"
                if w == "":
                    continue
                self.total_words = self.total_words + 1
                if w in self.stop:
                    self.stopped_count = self.stopped_count + 1
                    continue
                cnt = self.dictionary.get(w)
                if not cnt:
                    cnt = 0
                cnt = cnt + 1
                self.dictionary[w] = cnt
                self.words = self.words + 1


if __name__ == "__main__":
    # print( f"{sys.argv} {len(sys.argv)}")
    if "--help" in sys.argv:
        print("\nuse: \n\tpython3 build.py csv-filename ")
        # print("\tpython3 build.py directorname/^\n")
        print(
            '   A " starting a line means an italicized line, meaning a quote. It will also preserve a blank line.'
        )
        print(
            "   A | before and after a word means an italicized word i.e. a foreign word, ship name, etc."
        )
        print("   A | at the start of a line is a block quote, such as for song lyrics")
        print(
            "   Two |s at the start of a line is an escaped |, meaning the same as a | within a line"
        )
        print("\nQuotes cannot be escaped. Pipes are not legal characters in the text.")
        print(
            "\nThe input file must be a csv. Its first column is a directory name below the 'texts' directory. The second column is the filename. "
        )
        print(
            "\nThe file is a text file and should have a '.txt' extension; although, it doesn't have to. The first line of the text file is the title. There must be a blank line after the title and between paragraphs. A paragraph indentation will be added. Multiple blank lines are treated as one."
        )
        print(
            "\nThe output will be in the 'out' directory, named for the input csv file. Any existing output files will be overwritten at each run."
        )
        print(
            "\nIn the .csv file you can comment out a line by making the first character a '#'."
        )
        print("\n\n")
    else:
        print("For info: \n\tpython3 build.py --help \n")
        build = Build()

        import argparse

        parser = argparse.ArgumentParser()
        parser.add_argument(
            "book", help="a book. implemented as a config section.", nargs="?"
        )
        parser.add_argument("-config", help="a config .ini file", nargs="?")
        parser.add_argument(
            "-r", help="kill word and open output file", action="store_true"
        )

        args = parser.parse_args()

        if args.book:
            build.load_config(args.config, args.book)

        build.do()
        build.print_stats()

        print(f"args: {args}")
        if args.r:
            print("Killing Word")
            os.system("killall -9 Microsoft\ Word")

            print(f"open {build.OUT}")
            os.system(f"open {build.OUT}")


# ----------------------------------------
# not ultimately useful but interesting:
# nohup /Applications/Microsoft\ Word.app/Contents/MacOS/Microsoft\ Word /Users/davidkershaw/Desktop/desktop/new/space/builds/out/charles.docx > /dev/null &
